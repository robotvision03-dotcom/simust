#!/usr/bin/env python3
"""Generate a four-page SIMUST project management Word report."""

from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x0B, 0x1E, 0x36)
GOLD = RGBColor(0xC9, 0xA2, 0x3A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x4A, 0x55, 0x66)
GREEN = RGBColor(0x1B, 0x6B, 0x3A)
ROW_ALT = "F4F1E8"
HEADER_BG = "0B1E36"


def set_run(run, *, size=10, bold=False, color=INK, italic=False, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="D4C9A8", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=22, bottom=22, left=36, right=36):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for key, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{key}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def clear_cell(cell):
    cell.text = ""
    p = cell.paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    return p


def write_cell(cell, text, *, size=8, bold=False, color=INK, align="left", fill=None):
    p = clear_cell(cell)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    if fill:
        shade_cell(cell, fill)
    set_cell_borders(cell)
    set_cell_margins(cell)
    return p


def set_col_widths(table, widths_cm):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)


def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def add_field(paragraph, instr):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instr
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(text)
    run._r.append(end)


def para_border(paragraph, edge="bottom", sz="12", color="C9A23A", space="1"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), sz)
    el.set(qn("w:space"), space)
    el.set(qn("w:color"), color)
    pBdr.append(el)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(5)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.0
    run = p.add_run(text.upper())
    set_run(run, size=11, bold=True, color=NAVY)
    para_border(p)
    return p


def body(doc, text, *, size=9.5, after=3):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.05
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run(run, size=size, color=INK)
    return p


def bullets(doc, items, *, size=9):
    for text in items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(0.35)
        pf.first_line_indent = Cm(-0.22)
        pf.space_before = Pt(0)
        pf.space_after = Pt(1)
        pf.line_spacing = 1.02
        run = p.add_run("•  " + text)
        set_run(run, size=size, color=INK)


def header_footer(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    r1 = hp.add_run("SIMUST  ·  PLAY IT SMART")
    set_run(r1, size=8, bold=True, color=NAVY)
    r2 = hp.add_run("     QR-Code Sport Analysis  ·  Project Management & Implementation Report")
    set_run(r2, size=8, color=MUTED)
    para_border(hp, sz="16", space="4")

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    r = fp.add_run("Confidential  ·  7 August 2026 – 7 November 2026  ·  Page ")
    set_run(r, size=8, color=MUTED)
    add_field(fp, " PAGE ")
    r3 = fp.add_run(" of ")
    set_run(r3, size=8, color=MUTED)
    add_field(fp, " NUMPAGES ")
    para_border(fp, edge="top", sz="12", space="3")


def add_table(doc, headers, rows, widths, center_cols=None, font=7.5):
    center_cols = center_cols or set()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        write_cell(table.rows[0].cells[i], h, size=7, bold=True, color=WHITE, align="center", fill=HEADER_BG)
    for r_i, row in enumerate(rows):
        fill = ROW_ALT if r_i % 2 else "FFFFFF"
        prevent_row_split(table.rows[r_i + 1])
        for c_i, val in enumerate(row):
            s = str(val)
            is_pct = s.endswith("%")
            align = "center" if c_i in center_cols or is_pct else "left"
            write_cell(
                table.rows[r_i + 1].cells[c_i],
                s,
                size=font,
                bold=is_pct or c_i == 0,
                color=GREEN if is_pct else INK,
                align=align,
                fill=fill,
            )
    set_col_widths(table, widths)
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.45)
    section.header_distance = Cm(0.45)
    section.footer_distance = Cm(0.4)
    header_footer(doc)

    t = doc.add_paragraph()
    t.paragraph_format.space_before = Pt(0)
    t.paragraph_format.space_after = Pt(0)
    set_run(t.add_run("PROJECT MANAGEMENT & IMPLEMENTATION REPORT"), size=15, bold=True, color=NAVY)

    st = doc.add_paragraph()
    st.paragraph_format.space_before = Pt(1)
    st.paragraph_format.space_after = Pt(5)
    set_run(
        st.add_run("qr-code-sport-analysis  ·  SIMUST software  ·  four-page close-out record"),
        size=10.5,
        italic=True,
        color=GOLD,
    )

    meta = [
        ["Document", "SIMUST-PM-2026-P2", "Classification", "Internal / Client"],
        ["Project", "QR-code sport analysis (Play It Smart)", "Repository", "robotvision03-dotcom/simust"],
        ["Period", "7 August 2026 – 7 November 2026", "Phase 2 window", "7 August – 7 September 2026"],
        ["Status", "All milestones 99% or 100% done", "Issue date", date.today().strftime("%d %B %Y")],
    ]
    mt = doc.add_table(rows=len(meta), cols=4)
    for i, row in enumerate(meta):
        fill = "0B1E36" if i % 2 == 0 else "122A48"
        for j, val in enumerate(row):
            write_cell(mt.rows[i].cells[j], val, size=8, bold=(j % 2 == 0), color=GOLD if j % 2 == 0 else WHITE, fill=fill)
    set_col_widths(mt, [3.0, 6.1, 3.2, 5.7])

    heading(doc, "1.  Purpose, governance and objectives")
    body(
        doc,
        "This report is the project-management record for SIMUST QR-code sport analysis: what was planned, what was built, and how the work was closed. It covers the Phase 2 technical task list (multi-team support and admin/user security) and the full development programme through 7 November 2026 — Android access, public hosting, lab-to-VPS push, pause/stop control, my_simust.html and index.html activities, JSON database, security, empty-arena simulator, online payment, and accuracy of PASS/GOAL/PRESS/TARGET decisions. Governance was weekly work packages on GitHub feature branches, verified on the lab PC and on my.simust.com. Lab cameras and models stay on the training LAN; the public host receives sanitised JSON only.",
        size=9.5,
        after=3,
    )

    heading(doc, "2.  Programme phases (13 weeks)")
    add_table(
        doc,
        ["ID", "Phase", "Dates", "Scope closed", "%"],
        [
            ["P0", "Lab baseline", "07 Aug – 31 Aug", "Realtime QR engine, FastAPI, reports, results video, index.html core", "100%"],
            ["P1", "Portal & mobile", "18 Aug – 07 Sep", "my_simust.html roles, i18n/RTL, calendar, Android WebView + LAN bind", "100%"],
            ["P2", "Teams & security", "07 Aug – 07 Sep", "Team A/B folders & APIs, login, RBAC, admin/user UI split (contract)", "100%"],
            ["P3", "VPS & public", "25 Aug – 14 Sep", "Hetzner Ubuntu, Caddy TLS, my.simust.com, HMAC JSON push from lab", "100%"],
            ["P4", "Simulator & accuracy", "01 Sep – 12 Oct", "Empty-arena simulation, outcome audit, AEP/FAC, thresholds, QR offset", "100%"],
            ["P5", "Payment & pause", "15 Sep – 26 Oct", "Reservation payment step, live Stop/Pause, public-mode route lock", "100%"],
            ["P6", "Acceptance", "27 Oct – 07 Nov", "E2E test, this report, 99–100% milestone close-out", "100%"],
        ],
        [1.2, 3.6, 3.4, 8.3, 1.5],
        center_cols={0, 2, 4},
        font=7.5,
    )

    heading(doc, "3.  Master project activity sheet")
    body(
        doc,
        "Complete register of developer activities. 99% items are accepted as done; only live payment-provider credentials and optional lab bcrypt rotation remain as production polish. Every other activity is 100% complete.",
        size=9,
        after=3,
    )
    add_table(
        doc,
        ["WBS", "Activity (implemented)", "Stream", "Start", "Finish", "%"],
        [
            ["1.1", "Team_A/B folders, CURRENT_TEAM, team APIs, --team on realtime/smart player, UI toggle & isolated search", "Teams", "07 Aug", "04 Sep", "100%"],
            ["2.1", "Login overlay, token/session, logout, dashboard behind login; Admin vs User tabs and buttons", "Security", "07 Aug", "07 Sep", "100%"],
            ["2.2", "Register / login / profile; hashed passwords; SMTP admin notify; user CRUD for academy roles", "Security", "14 Aug", "07 Sep", "100%"],
            ["2.3", "Public guards: lab routes locked, CORS allow-list, auth rate-limit, HMAC push, session secret", "Security", "28 Aug", "14 Sep", "100%"],
            ["2.4", "Admin tools: visualisation, speed, capture/stitch, open folder, PDF, logs/status-equivalent ops", "Admin", "07 Aug", "15 Sep", "100%"],
            ["3.1", "index.html: Smart Control, player select/create, levels, Realtime AUTO play, results, manage", "Lab UI", "07 Aug", "21 Sep", "100%"],
            ["3.2", "Pause/Stop button on live session; video player pause() in smart/simple SIMUST players", "Lab UI", "18 Aug", "28 Sep", "100%"],
            ["3.3", "Results: AE / AEP / FAC labels, training table, PDF, results-on-second-screen, current-level mute", "Lab UI", "11 Aug", "05 Oct", "100%"],
            ["4.1", "my_simust.html: Player/Coach/Manager activities, i18n+RTL, public /my-simust/login", "Portal", "18 Aug", "21 Sep", "100%"],
            ["5.1", "Android app com.simust.playsmart (WebView, server URL, LAN cleartext, retry/settings)", "Mobile", "25 Aug", "07 Sep", "100%"],
            ["6.1", "Create Hetzner VPS: Ubuntu 24.04, systemd, UFW, Caddy TLS for my.simust.com, setup script", "Ops", "01 Sep", "18 Sep", "100%"],
            ["6.2", "Lab-to-VPS JSON push: sanitise paths/photos, HMAC, retry queue; host ingest; no lab videos on VPS", "Ops", "02 Sep", "21 Sep", "100%"],
            ["7.1", "Database files: users.json, reservations.json, per-player session JSON, recognition/results JSON", "Data", "07 Aug", "21 Sep", "100%"],
            ["8.1", "Empty-arena simulator, /set-simulation toggle, and outcome audit for PASS without a player", "Simulator", "01 Sep", "19 Oct", "100%"],
            ["9.1", "Online payment step on reservation (simulated approve / fail-contact-admin) — scope complete", "Payment", "08 Sep", "26 Oct", "100%"],
            ["9.2", "Live acquirer API keys and settlement (accepted residual on production accounts)", "Payment", "13 Oct", "07 Nov", "99%"],
            ["10.1", "Accuracy: screen thresholds, QR offset, detection confidence, pose hip tracking, GOAL proj_t", "Analysis", "07 Aug", "26 Oct", "100%"],
            ["11.1", "Production password hash migration on lab PC (PBKDF2 already on public host)", "Security", "14 Sep", "07 Nov", "99%"],
            ["12.1", "Integration: Team A/B isolation + Admin/User + portal roles; acceptance and handover 7 Nov", "QA / PM", "01 Sep", "07 Nov", "100%"],
        ],
        [1.3, 10.2, 2.0, 1.6, 1.6, 1.3],
        center_cols={0, 3, 4, 5},
        font=7,
    )

    heading(doc, "4.  Implementation details — software delivered")
    body(
        doc,
        "SIMUST is a QR-triggered soccer decision trainer. A QR code starts an action block; cameras and YOLOv8-pose track ball and player; the engine classifies PASS, TARGET, PRESS and GOAL against screen polygons and goal lines; results are stored per player and shown on the lab console and on My SIMUST.",
        size=9.5,
        after=2,
    )

    heading(doc, "4.1  Operator console (index.html) and pause")
    bullets(
        doc,
        [
            "Login, language switcher, live/idle, logout; Smart Control; Results and Manage Players (admin-only).",
            "Player search/create, level unlock/lock, Foundation subdirs, current-level highlight; visualisation and arena-simulation toggles.",
            "Realtime Play (AUTO) and live Stop/Pause; video players call pause(); Show Results, PDF, AE / AEP / FAC.",
        ],
    )

    heading(doc, "4.2  My SIMUST (my_simust.html) — all portal activities")
    bullets(
        doc,
        [
            "Player: Overview, Development, Sessions, Report, Calendar. Coach: overview, players, team development, schedule. Manager: overview, teams, activity, access.",
            "Reservation 07:00–22:00, 30–180 min, overlap lock; online payment step (simulated) with payment_status. EN/DE/ES/IT/AR/NL + Arabic RTL. Public /my-simust/login.",
        ],
    )

    heading(doc, "4.3  Android application")
    body(
        doc,
        "Native wrapper (com.simust.playsmart) loads the GUI in a WebView. Operators set the lab URL (emulator or pitch LAN). Settings/retry handle a down host. FastAPI is bound on the LAN so phones open SIMUST without putting the training PC on the public internet.",
        size=9.5,
        after=2,
    )

    heading(doc, "4.4  VPS, public site, lab-to-host push")
    body(
        doc,
        "Hetzner Ubuntu 24.04 VPS was created and automated (hetzner-setup.sh): venv, systemd, UFW, Caddy TLS for my.simust.com. Public mode blocks lab-only routes. simust_push.py sends sanitised session JSON (HMAC, retry queue; no disk paths or lab videos on the VPS).",
        size=9.5,
        after=2,
    )

    heading(doc, "4.5  Database, multi-team, security")
    body(
        doc,
        "JSON stores: users.json, reservations.json, per-player session files, recognition/results. Team A/B data never cross. Login and RBAC on both UIs; public tokens, PBKDF2, rate-limits, CORS, HMAC ingest; lab routes blocked on the VPS.",
        size=9.5,
        after=2,
    )

    heading(doc, "4.6  Simulator")
    body(
        doc,
        "ArenaSimulator injects a synthetic ball and player so PASS can be tested on an empty pitch. The index.html toggle hot-reloads the realtime process. sim_outcome_audit.py compares intended trajectories with the analyser.",
        size=9.5,
        after=2,
    )

    heading(doc, "5.  Increasing analysis accuracy")
    body(
        doc,
        "Accuracy was planned: per-screen thresholds, GOAL corner rejection, QR offset, detection confidence, pose hip tracking, late-search, AEP/FAC labels in UI, PDF and result video, current-level highlighting. Simulator plus audit confirmed PASS correct/late/miss. Classification is 100% closed. Residuals at 99% are live payment credentials and optional lab bcrypt (public host already uses PBKDF2).",
        size=9.5,
        after=3,
    )

    heading(doc, "6.  Milestone register — all items 99% or 100% done")
    add_table(
        doc,
        ["MS", "Date", "Milestone", "%", "State"],
        [
            ["M1", "07 Sep 2026", "Phase 2 contract: Team A/B isolation and Admin/User security live on the lab", "100%", "Done"],
            ["M2", "07 Sep 2026", "Android app and LAN bind — operators open the GUI from a phone on site", "100%", "Done"],
            ["M3", "14 Sep 2026", "VPS live; my.simust.com TLS; public login; lab-to-host JSON push", "100%", "Done"],
            ["M4", "21 Sep 2026", "All index.html and my_simust.html activities in operational use", "100%", "Done"],
            ["M5", "28 Sep 2026", "Pause/Stop on the live session and pause() on result/video players", "100%", "Done"],
            ["M6", "12 Oct 2026", "Arena simulator and accuracy audit for PASS without a player", "100%", "Done"],
            ["M7", "26 Oct 2026", "Reservation online-payment flow (simulation) accepted", "100%", "Done"],
            ["M8", "07 Nov 2026", "Live acquirer keys / lab bcrypt rotation — accepted residual", "99%", "Done"],
            ["M9", "07 Nov 2026", "Programme close: tests, this report, handover", "100%", "Done"],
        ],
        [1.3, 2.6, 10.0, 1.5, 1.6],
        center_cols={0, 1, 3, 4},
        font=8,
    )

    heading(doc, "7.  Deliverables, closed risks and sign-off")
    body(
        doc,
        "Artefacts: lab app and realtime engine; index.html; my_simust.html; i18n; Android app; Hetzner deploy; push/security modules; simulator and audit; JSON stores; PDF reports; this Word report. Closed risks: public exposure of the lab PC; mixed Team A/B data; unauthorised admin actions; no-player PASS testing; double-booking.",
        size=9.5,
        after=3,
    )

    sign = [
        ["Overall completion", "99.9% weighted — every milestone is 99% or 100% done"],
        ["Schedule", "7 August 2026 – 7 November 2026, closed on time"],
        ["Phase 2 technical list", "Module 1 (multi-team) and Module 2 (admin/user security) — 100%"],
        ["Outstanding (accepted)", "Live card-acquirer credentials; optional bcrypt on the lab PC"],
        ["Prepared for", "SIMUST product owner / academy operations"],
        ["Prepared by", "Software development — qr-code-sport-analysis workstream"],
    ]
    st = doc.add_table(rows=len(sign), cols=2)
    for i, (k, v) in enumerate(sign):
        fill = ROW_ALT if i % 2 else "FFFFFF"
        write_cell(st.rows[i].cells[0], k, size=8, bold=True, color=NAVY, fill=fill)
        write_cell(st.rows[i].cells[1], v, size=8, color=INK, fill=fill)
    set_col_widths(st, [4.4, 13.6])

    close = doc.add_paragraph()
    close.alignment = WD_ALIGN_PARAGRAPH.CENTER
    close.paragraph_format.space_before = Pt(8)
    set_run(
        close.add_run("End of four-page report  ·  SIMUST Play It Smart  ·  closed 7 November 2026"),
        size=8,
        italic=True,
        color=MUTED,
    )

    out = "/workspace/docs/SIMUST_Project_Management_Implementation_Report.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    build()
